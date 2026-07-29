#!/usr/bin/env python3
"""
ROCF 电子测评系统 — 办公自动化脚本
=====================================
功能：
  - 读取导出的 JSON / CSV 数据文件
  - 批量生成 PDF 报告
  - 将数据写入 Word 文档（python-docx）
  - 命令行入口

用法：
  python office_auto.py report --input data.json --output report.pdf
  python office_auto.py to-word --input data.csv --output summary.docx
  python office_auto.py batch-report --glob "exports/*.json" --out-dir reports/
"""

import argparse
import csv
import json
import os
import sys
from datetime import datetime
from glob import glob
from pathlib import Path


# ---------------------------------------------------------------------------
# 数据读取
# ---------------------------------------------------------------------------

def load_json(path: str) -> dict:
    """读取单个 JSON 数据文件。"""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_csv(path: str) -> list[dict]:
    """读取 CSV 数据文件，返回字典列表。"""
    rows = []
    with open(path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def find_files(pattern: str) -> list[str]:
    """glob 查找匹配的文件，按名称排序。"""
    return sorted(glob(pattern, recursive=False))


# ---------------------------------------------------------------------------
# 评分摘要
# ---------------------------------------------------------------------------

def build_summary(data: dict | list[dict]) -> dict:
    """从原始数据中提取评分摘要信息。"""
    if isinstance(data, list):
        records = data
    else:
        records = [data]

    total = len(records)
    scores = []
    for r in records:
        score = _extract_score(r)
        if score is not None:
            scores.append(score)

    avg = sum(scores) / len(scores) if scores else 0
    return {
        "record_count": total,
        "score_count": len(scores),
        "average_score": round(avg, 2),
        "min_score": min(scores) if scores else None,
        "max_score": max(scores) if scores else None,
    }


def _extract_score(record: dict) -> float | None:
    """尝试从一条记录中提取数值评分。"""
    for key in ("score", "total_score", "raw_score", "总分", "得分"):
        val = record.get(key)
        if val is not None:
            try:
                return float(val)
            except (ValueError, TypeError):
                pass
    return None


# ---------------------------------------------------------------------------
# PDF 报告生成（纯文本型，依赖 reportlab 可选）
# ---------------------------------------------------------------------------

def generate_pdf_report(data: dict, output_path: str):
    """生成 PDF 报告。优先使用 reportlab，回退到纯文本。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
        _write_pdf_reportlab(data, output_path)
    except ImportError:
        print("[WARN] reportlab 未安装，回退为纯文本 .txt 报告")
        txt_path = output_path.rsplit(".", 1)[0] + ".txt"
        _write_text_report(data, txt_path)


def _write_pdf_reportlab(data: dict, output_path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    y = height - 30 * mm
    c.setFont("Helvetica-Bold", 18)
    c.drawString(20 * mm, y, "ROCF Rey-Osterrieth Complex Figure Test")
    y -= 12 * mm
    c.setFont("Helvetica", 11)
    c.drawString(20 * mm, y, f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    y -= 10 * mm

    summary = data.get("_summary", build_summary(data))
    lines = [
        f"Records: {summary.get('record_count', 0)}",
        f"Scored:  {summary.get('score_count', 0)}",
        f"Average: {summary.get('average_score', 'N/A')}",
        f"Range:   {summary.get('min_score', 'N/A')} - {summary.get('max_score', 'N/A')}",
    ]
    c.setFont("Helvetica", 11)
    for line in lines:
        c.drawString(20 * mm, y, line)
        y -= 6 * mm

    y -= 6 * mm
    c.setFont("Helvetica", 9)
    c.drawString(20 * mm, y, "Scoring method: Osterrieth 36-point system")
    c.save()
    print(f"[OK] PDF saved → {output_path}")


def _write_text_report(data: dict, output_path: str):
    summary = data.get("_summary", build_summary(data))
    lines = [
        "ROCF Rey-Osterrieth Complex Figure Test — Report",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        f"Records: {summary.get('record_count', 0)}",
        f"Scored:  {summary.get('score_count', 0)}",
        f"Average: {summary.get('average_score', 'N/A')}",
        f"Range:   {summary.get('min_score', 'N/A')} - {summary.get('max_score', 'N/A')}",
        "",
        "Scoring: Osterrieth 36-point system",
    ]
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[OK] Text report saved → {output_path}")


# ---------------------------------------------------------------------------
# Word 文档生成
# ---------------------------------------------------------------------------

def generate_word(data: dict | list[dict], output_path: str):
    """将数据写入 Word 文档。"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx 未安装，请执行: pip install python-docx")
        sys.exit(1)

    if isinstance(data, list):
        records = data
    else:
        records = [data]

    doc = Document()
    doc.add_heading("ROCF 电子测评系统 — 数据报告", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"记录数量：{len(records)}")

    summary = build_summary(records)
    doc.add_heading("评分摘要", level=2)
    table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    items = [
        ("平均分", summary["average_score"]),
        ("最低分", summary["min_score"]),
        ("最高分", summary["max_score"]),
        ("已评分", summary["score_count"]),
        ("总记录", summary["record_count"]),
    ]
    for i, (k, v) in enumerate(items):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)

    doc.add_heading("详细记录", level=2)
    for idx, record in enumerate(records, 1):
        doc.add_heading(f"记录 #{idx}", level=3)
        for k, v in record.items():
            doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    doc.save(output_path)
    print(f"[OK] Word saved → {output_path}")


# ---------------------------------------------------------------------------
# 命令行入口
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="ROCF 电子测评系统 — 办公自动化工具",
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # report
    p_report = sub.add_parser("report", help="生成单份 PDF 报告")
    p_report.add_argument("--input", required=True, help="输入的 JSON/CSV 文件路径")
    p_report.add_argument("--output", default="rocf_report.pdf", help="输出 PDF 路径")

    # batch-report
    p_batch = sub.add_parser("batch-report", help="批量生成 PDF 报告")
    p_batch.add_argument("--glob", required=True, dest="glob_pattern", help="文件匹配模式，如 'exports/*.json'")
    p_batch.add_argument("--out-dir", default="reports", help="输出目录，默认 reports/")

    # to-word
    p_word = sub.add_parser("to-word", help="将数据写入 Word 文档")
    p_word.add_argument("--input", required=True, help="输入的 JSON/CSV 文件路径")
    p_word.add_argument("--output", default="rocf_summary.docx", help="输出 .docx 路径")

    args = parser.parse_args()

    if args.command == "report":
        data = _load_any(args.input)
        generate_pdf_report(data, args.output)

    elif args.command == "batch-report":
        files = find_files(args.glob_pattern)
        if not files:
            print(f"[WARN] 未找到匹配文件: {args.glob_pattern}")
            sys.exit(1)
        os.makedirs(args.out_dir, exist_ok=True)
        for f in files:
            data = _load_any(f)
            stem = Path(f).stem
            out_path = os.path.join(args.out_dir, f"{stem}_report.pdf")
            generate_pdf_report(data, out_path)
        print(f"[OK] 批量完成，共 {len(files)} 份报告 → {args.out_dir}")

    elif args.command == "to-word":
        data = _load_any(args.input)
        generate_word(data, args.output)


def _load_any(path: str):
    """根据扩展名自动选择加载方式。"""
    ext = Path(path).suffix.lower()
    if ext == ".json":
        return load_json(path)
    elif ext == ".csv":
        return load_csv(path)
    else:
        print(f"[ERROR] 不支持的文件格式: {ext}")
        sys.exit(1)


if __name__ == "__main__":
    main()
